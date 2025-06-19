"""
For app.py integration
"""
import os
import re
import openai

from io import BytesIO
from docxcompose.composer import Composer

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import streamlit as st


OPENAI_API_KEY = st.secrets["openai_api_key"]
CORRECT_PASSWORD = st.secrets["app_password"]

# Set OpenAI key
openai.api_key = OPENAI_API_KEY

# MODEL = "gpt-4o-mini"
MODEL = "gpt-4o"

def generate_one_pager(company_name, content_dict, output_path) -> BytesIO:
    doc = Document("Template.docx")
    
    # Set to portrait and A4
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)

    # Set standard margins (optional tweak)
    section.top_margin = section.bottom_margin = Pt(72)  # 1 inch
    section.left_margin = section.right_margin = Pt(72)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)

    # Set global line spacing to 1.3
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = 1.15

    # Add each section
    for heading, text in content_dict.items():
        # Heading
        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run(heading)
        heading_run.font.name = 'Calibri'
        heading_run.font.size = Pt(18)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(255, 153, 0)

        # Content
        # Add quotes for Vision and Mission Statements
        if heading == "Vision Statement" or heading == "Mission Statement":
            text = "“" + text + "”"  

        body_para = doc.add_paragraph(text)
        for run in body_para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(12)

        # Double New Lines between Paragraph and New Heading (except at the very last heading)
        if heading != "Definition of Success":
            doc.add_paragraph()
            doc.add_paragraph()

    # Save
    # doc.save(output_path)
    doc_cover = Document()
    section = doc_cover.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)

    # Set standard margins (optional tweak)
    section.top_margin = section.bottom_margin = Pt(72)  # 1 inch
    section.left_margin = section.right_margin = Pt(72)

    style = doc_cover.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)

    # Set global line spacing to 1.3
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = 1.15

    insert_cover_page(doc_cover, company_name=company_name, logo_path="Logo3.png")
    composer = Composer(doc_cover)

    composer.append(doc)

    # doc_cover.save(output_path)
    buffer = BytesIO()
    doc_cover.save(buffer)
    buffer.seek(0)  # Move back to the beginning so Streamlit can read it
    return buffer

def build_prompt(minutes, company_name):
    combined_prompt_template = f"""
    You are a professional business strategist who has just run a workshop for a business called "{company_name}"

    All writing should use British English spelling and conventions. Where appropriate, expand upon the ideas captured during the workshop to ensure clarity, completeness, and usefulness.

    Write this in a professional tone using clear, direct language. Avoid overly formal or common ChatGPT phrases like "delve," "poise," "robust," etc.

    You are helping summarize a business strategy workshop. You do not need to create a title, as we have a cover page already made.

    You are to use the workshop minutes, provided below, to generate a concise, high-level, compelling one-page summary document. You will be summarising 6 six secitons, each as a succinct paragraph in 35 words or less. The sections are:

     - **Vision Statement** Write a single sentence that communicates a clear and inspiring long-term vision for the organization.

     - **Mission Statement** Write a compelling mission statement.

     - **Customers** Summarize the key customers discussed, written as a succinct paragraph.

     - **Value Proposition** Generate a clear, concise, and non-repetitive value proposition statement.

     - **Products and Services** Write a brief, clear description of the organization's core products and services.

     - **Definition of Success** Define what success looks like for the organization based on the workshop discussion.

    --- WORKSHOP MINUTES START ---

    {minutes}

    --- WORKSHOP MINUTES END ---
    """
    return combined_prompt_template

def generate_combined_summary(minutes, company_name):
    """Generates the entire one-pager using the combined prompt."""
    prompt = build_prompt(minutes, company_name)
    response = openai.ChatCompletion.create(
        model=MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7
    )
    return response['choices'][0]['message']['content'].strip()

def insert_cover_page(doc, company_name, logo_path=None):
    # Add blank lines to push text down
    for _ in range(10):  # Adjust number as needed for vertical spacing
        doc.add_paragraph()

    # Add Company Name (centered, large, orange, bold)
    para1 = doc.add_paragraph()
    run1 = para1.add_run(company_name)
    run1.font.name = 'Calibri'
    run1.font.size = Pt(44)
    run1.font.color.rgb = RGBColor(255, 153, 0)  # Orange (#FF9900)
    run1.bold = True
    para1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add "Strategy Report" below
    para2 = doc.add_paragraph()
    run2 = para2.add_run("1-Page Strategy")
    run2.font.name = 'Calibri'
    run2.font.size = Pt(44)
    run2.font.color.rgb = RGBColor(255, 153, 0)  # Orange (#FF9900)
    run2.bold = True
    para2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Optional spacing before logo
    doc.add_paragraph()

    # Insert logo if provided
    if logo_path:
        logo_para = doc.add_paragraph()
        logo_run = logo_para.add_run()
        logo_run.add_picture(logo_path, width=Inches(2))
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add a page break after the cover page
    doc.add_page_break()

def split_one_pager_sections(text: str) -> dict:
    """
    Extract sections from a one-pager AI response using bold headings (e.g. **Vision Statement**).
    Returns a dictionary with section names as keys and content as values.
    """
    # Regex to split by bold headings (e.g. **Vision Statement**)
    pattern = r"\*\*(.*?)\*\*"

    # Find all section titles
    matches = list(re.finditer(pattern, text))
    content_dict = {}

    for i, match in enumerate(matches):
        heading = match.group(1).strip()

        # Start and end of the content chunk
        start = match.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)

        # Extract content, clean extra spaces/newlines
        section_text = text[start:end].strip()
        section_text = re.sub(r"\n+", " ", section_text).strip()
        content_dict[heading] = section_text

    return content_dict


def generate_one_pager_docx(minutes, filename, company_name) -> BytesIO:
    one_pager_text = generate_combined_summary(minutes, company_name)
    content_dict  = split_one_pager_sections(one_pager_text)
    buffer = generate_one_pager(company_name, content_dict, filename)

    return buffer
