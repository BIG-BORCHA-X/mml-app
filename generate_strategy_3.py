"""
This update:
 - Adds new sections dynamically
   - ***Section Heading***
   - ^^ Looks for this in minutes and does the rest automatically.
   - Places new secitons either before BM sections, after before Recommendations,
     based on if the new headings were found before or after the BM sections.

Should be the last major update before tweaking starts
"""

import os
import json
import re
import openai

from dotenv import load_dotenv
from datetime import datetime       # for file signature
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import streamlit as st

# Load API key from .env
# load_dotenv()
# openai.api_key = os.getenv("OPENAI_API_KEY")

OPENAI_API_KEY = st.secrets["openai_api_key"]
CORRECT_PASSWORD = st.secrets["app_password"]

# Set OpenAI key
openai.api_key = OPENAI_API_KEY

# ----------- Config -----------
MODEL = "gpt-4o"
# MODEL = "gpt-4o-mini"
MAX_TOKENS = 800                                    # Per section - roughly 600-700 words MAX

# Sections to generate
TESTING = False
if TESTING:
    # tweaking
    SECTIONS = [
        ["Our Approach", 250],
        ["Scope of Project", 300],
        ["Vision", 150+75],
        ["Conclusion", 150+50]
    ]
else:
    # fr fr
    SECTIONS = [
        ["Our Approach", 250],
        ["Scope of Project", 300],
        ["Definition of Success", 420],
        ["Purpose of Starting the Business", 150+750],
        ["Vision", 150+75],
        ["Mission", 200+50],
        ["Goals", 200+25],
        ["Product Service Offering", 150+75],
        # Business Model Segments until Recommendations
        ["Customer Segments", 150+75],
        ["Value Proposition", 200+50],
        ["Channels", 150+75],
        ["Customer Relationships", 150+75],
        ["Revenue Streams", 150+75],   # Smallest section in Sample
        ["Key Resources", 200+50],
        ["Key Activities", 200+50],
        ["Key Partners", 200+50],
        ["Cost Structure", 175+50],
        ["Recommendations", 600+50],
        ["Conclusion", 150+50]
    ]

BM_SECTIONS = ["Customer Segments", "Value Proposition", "Channels", "Customer Relationships",
               "Revenue Streams", "Key Resources", "Key Activities", "Key Partners", "Cost Structure"]

# ----------- Functions -----------
def clean_heading(heading):
    # Trim leading/trailing whitespace
    heading = heading.strip()

    # Remove leading number/dot patterns, with or without space (e.g., "2. ", "2.1.", "10.2. ")
    heading = re.sub(r'^\d+(?:\.\d+)*\.?\s*', '', heading)

    # Remove trailing colon ":  "
    heading = heading.rstrip(':')

    # Final strip to catch " :  "
    heading = heading.strip()

    # Smart capitalize
    return smart_capitalize(heading)

def insert_new_sections_and_prompts(SECTIONS: list, prompts, before_bsm, after_bsm, default_token_limit=300):
    # prompts_list = list(prompts.values()) # No longer a dict, parsed as a list

    updated_sections = SECTIONS.copy()
    updated_prompts = prompts.copy()
    
    # Find anchor points
    insert_after_product = next((i for i, (h, _) in enumerate(updated_sections) if h == "Product Service Offering"), None)
    insert_after_cost = next((i for i, (h, _) in enumerate(updated_sections) if h == "Cost Structure"), None)

    if insert_after_product is None or insert_after_cost is None:
        raise ValueError("Required anchor headings not found in SECTIONS.")
    
    # === Insert before_bsm ===
    for offset, heading in enumerate(before_bsm):
        heading = clean_heading(heading)
        insert_index = insert_after_product + 1 + offset
        updated_sections.insert(insert_index, (heading, default_token_limit))
        updated_prompts.insert(insert_index, generate_new_section_prompt(heading))

    # Adjust cost index if before_bsm added elements
    insert_after_cost += len(before_bsm)

    # === Insert after_bsm ===
    for offset, heading in enumerate(after_bsm):
        heading = clean_heading(heading)
        insert_index = insert_after_cost + 1 + offset
        updated_sections.insert(insert_index, (heading, default_token_limit))
        updated_prompts.insert(insert_index, generate_new_section_prompt(heading))

    return updated_sections, updated_prompts

def find_section_position(minutes: str, heading: str, anchor_phrase="Business Structure Mapping") -> str:
    # Find position of the anchor
    anchor_index = minutes.lower().find(anchor_phrase.lower())
    
    # Find position of the heading in the minutes (in ***Heading*** format)
    heading_pattern = f"***{heading}***"
    heading_index = minutes.lower().find(heading_pattern.lower())
    
    if heading_index == -1:
        return "unknown"
    if anchor_index == -1:
        return "no anchor"
    
    return "before" if heading_index < anchor_index else "after"

def smart_capitalize(text):
    result = ''
    capitalize_next = True
    for char in text:
        if capitalize_next and char.isalpha():
            result += char.upper()
            capitalize_next = False
        else:
            result += char
        if char == ' ':
            capitalize_next = True
    return result

def find_new_headings(minutes):
    # Locates headings like: ***New Heading***
    matches = re.findall(r'\*\*\*(.*?)\*\*\*', minutes)

    # Clean and title-case each match
    new_headings = [smart_capitalize(match) for match in matches]

    return new_headings

def generate_new_section_prompt(heading: str) -> str:
    return f"""Write the content for the '{heading}' section for the strategy report.

Start with a short paragraph introducing the topic: {heading}. You may interpret what this section likely covers based on the heading. End the paragraph with a sentence that introduces the bullet points (e.g. 'Below are our key points:')

Then create a bullet-pointed list. Each bullet point should:
- Start with a hyphen and a bolded heading like `- **Example Bullet Point**`
- Be concise, and no more than one full sentence
- Optionally add a concise explanation on the same line as the bolded text like `- **Example Bullet Point:** Example concise explanation`
- Avoid blank lines between bullets

Finish with a short paragraph that reflects on the importance of this section's content in the overall strategy or operations of the business.

Use a professional and helpful tone. This section will be part of a formal business report summarizing a strategy workshop."""

def save_raw_text(section_name: str, content: str, output_dir="debug_outputs"):
    os.makedirs(output_dir, exist_ok=True)
    filename = os.path.join(output_dir, f"{section_name}.txt")
    
    with open(filename, "w", encoding="utf-8") as f:
        f.write(content)

def load_section_from_file(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()
    return re.sub(r"\n{2,}", "\n", content.strip())

# Load prompt library from JSON file
def load_prompt_library(filepath):
    with open(filepath, "r", encoding="utf-8") as file:
        return json.load(file)

def build_global(company_name):

    GLOBAL_PROMPT = f"""
    You are a professional business strategist who has just run a workshop for a business called "{company_name}". You've gathered key insights that now need to be turned into a professional, high-quality report.

    The objective is to generate a well-written, detailed, and structured section of a business strategy report based on the provided workshop minutes. The writing must be clear, actionable, and appropriate for a professional audience.

    All writing should use British English spelling and conventions. Where appropriate, expand upon the ideas captured during the workshop to ensure clarity, completeness, and usefulness.
    
    Write this in a professional tone using clear, direct language. Avoid overly formal or common ChatGPT phrases like "delve," "poise," "robust," etc.
    """
    return GLOBAL_PROMPT

# Build the full prompt for a section
def build_prompt(global_prompt, minutes, section_prompt, token_limit):
    section_prompt += "\n\nDo not include a section heading at the start of your response."

    return (
        f"{global_prompt}\n\n"
        f"=== Workshop Minutes ===\n{minutes}\n\n"
        f"=== Section Instructions ===\n{section_prompt}\n\n"
        f"Please limit your response to approximately {token_limit} tokens or fewer."
    )

def normalize_newlines(text: str) -> str:
    """
    Replace multiple consecutive newlines with a single newline.
    """
    return re.sub(r'\n{2,}', '\n', text.strip())

def generate_static_approach_section(company_name):
    # Section is "Cookie Cutter", indentical each time except client name.
    content = f"\nMomentum Mind Lab engaged with you to evaluate the current position of {company_name} and develop a comprehensive organisational model and process for taking this forward. We embraced a customer-centred approach to developing solutions following the principles of Design Thinking (DT). We started the process by discovering your goals, expectations, strengths and capabilities. This allowed us to assess what is moving the business forward and what is holding it back, subsequently acknowledging the need to focus on specific aspects of the business in consideration of the goals and capabilities of {company_name}.\n\nAs part of the definition process, we mapped the organisation's structural model to gain clarity about the different elements of the organisation. This entailed defining why the business was started, what the product is as well as who it was created for. This provided a foundation for a macro-level organisational process mapping for identifying the specific areas of the organisation that need to be prioritised to increase efficiency. As a result, key areas of focus were defined, and a clear and detailed strategic action plan was developed for you, which indicates what actions need to be taken, what are the tasks associated with each action, and success criteria to monitor your progress."
    return content

def generate_static_scope_section(company_name):
    plural_company = f"{company_name}'s"
    quoted_company = f'"{plural_company}"'
    content = f"\nDear {company_name},\n\nThank you for giving us the opportunity to work with you during this workshop. Your enthusiastic and committed participation in the workshop was instrumental in shaping this report. Your dedication to {quoted_company} mission and your willingness to engage in collaborative strategic planning has been truly inspiring.\n"
    return content

def insert_cover_page(doc, company_name, logo_path=None):
    # Add blank lines to push text down
    for _ in range(4):  # Adjust number as needed for vertical spacing
        doc.add_paragraph()

    # Add Company Name (centered, large, orange, bold)
    para1 = doc.add_paragraph()
    run1 = para1.add_run(company_name)
    run1.font.name = 'Calibri'
    run1.font.size = Pt(44)
    run1.font.color.rgb = RGBColor(255, 153, 0)  # Orange (#FF9900)
    run1.bold = True
    para1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add "Strategy Report" below
    para2 = doc.add_paragraph()
    run2 = para2.add_run("Strategy Report")
    run2.font.name = 'Calibri'
    run2.font.size = Pt(44)
    run2.font.color.rgb = RGBColor(255, 153, 0)  # Orange (#FF9900)
    run2.bold = True
    para2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Optional spacing before logo
    doc.add_paragraph()

    # Insert logo if provided
    if logo_path:
        logo_para = doc.add_paragraph()
        logo_run = logo_para.add_run()
        logo_run.add_picture(logo_path, width=Inches(2))
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add a page break after the cover page
    doc.add_page_break()

def extract_company_name(minutes, model="gpt-4o-mini"):
    prompt = (
        "Extract the name of the company or client mentioned in the following workshop minutes.\n"
        "Only return the company name. No explanation, no punctuation.\n\n"
        f"{minutes}"
    )

    response = openai.ChatCompletion.create(
        model=model,
        messages=[{"role": "user", "content": prompt}],
        max_tokens=20,
        temperature=0
    )
    return response["choices"][0]["message"]["content"].strip()

def insert_table_of_contents(doc):
    """
     - Does not work, XML field codes don't work.
     - Table of Contents not accessible
     - Can't even generate a Blank or Empty ToC to be manually update

     - Solution: Insert a blank page and manually insert ToC and Update it.
    
    """
    # paragraph = doc.add_paragraph()
    # run = paragraph.add_run()

    # fldChar1 = OxmlElement('w:fldChar')
    # fldChar1.set(qn('w:fldCharType'), 'begin')

    # instrText = OxmlElement('w:instrText')
    # instrText.set(qn('xml:space'), 'preserve')
    # instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

    # fldChar2 = OxmlElement('w:fldChar')
    # fldChar2.set(qn('w:fldCharType'), 'separate')

    # fldChar3 = OxmlElement('w:fldChar')
    # fldChar3.set(qn('w:fldCharType'), 'end')

    # r_element = run._r
    # r_element.append(fldChar1)
    # r_element.append(instrText)
    # r_element.append(fldChar2)
    # r_element.append(fldChar3)

    doc.add_paragraph()  # Optional spacing
    doc.add_page_break()

def is_bullet_point(line):
    stripped = line.strip()
    return bool(re.match(r"^[-–—•●]\s+", stripped))

def add_markdown_bold_paragraph(doc, text, style="Normal"):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(0)

    # Indent bullets only
    if style == "List Bullet":
        paragraph.paragraph_format.left_indent = Inches(0.5)

    # Split into parts by bold markers (**...**)
    parts = re.split(r"(\*\*.*?\*\*)", text)

    for part in parts:
        run = paragraph.add_run()
        run.font.name = 'Calibri'
        run.font.size = Pt(12)

        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
        else:
            run.text = part

    return paragraph

def insert_logo(doc, image_path, width_in_inches=2):
    if image_path:
        para = doc.add_paragraph()
        run = para.add_run()
        run.add_picture(image_path, width=Inches(width_in_inches))
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Helper function: add landscape section break
def set_landscape(document):
    section = document.sections[-1]
    
    # Set A4 size
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.orientation = WD_ORIENT.LANDSCAPE

# Main writing function
def write_to_docx(file_path, global_prompt, minutes, prompt_library, sections, company_name, status_area=None) -> BytesIO:
    doc = Document()
    set_landscape(doc)

    # Set normal margins
    section = doc.sections[-1]
    inch = Inches(1)
    section.top_margin = inch
    section.bottom_margin = inch
    section.left_margin = inch
    section.right_margin = inch

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)

    # Set global line spacing to 1.3
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = 1.3

    # company_name = extract_company_name(minutes)
    insert_cover_page(doc, company_name=company_name, logo_path="Logo3.png")
    # Currently jsut a blank page
    insert_table_of_contents(doc)

    # Prompts
    prompt_values = list(prompt_library.values())

    # Track whether we've already added the "Business Model" heading
    inserted_bm_heading = False

    # Check for new headings / sections
    new_sections = find_new_headings(minutes)
    new_section_limit = 300                     # tokens

    # For rough positioning in document
    before_bm = []
    after_bm = []

    for heading in new_sections:
        pos = find_section_position(minutes, heading)
        if pos == "before":
            before_bm.append(heading)
        elif pos == "after":
            after_bm.append(heading)

    updated_sections, updated_prompts = insert_new_sections_and_prompts(sections, prompt_values, before_bm, after_bm)
    sections = updated_sections
    prompt_values = updated_prompts

    for i, (heading, token_limit) in enumerate(sections):
        if status_area:
            status_area.text(f"Generating {heading}...")
        else:
            print(f"Generating section: {heading}...")
        
        if heading == "Our Approach":
            content = generate_static_approach_section(company_name)
        elif heading == "Scope of Project":
            section_prompt = prompt_values[i]
            full_prompt = build_prompt(global_prompt, minutes, section_prompt, token_limit)
            static_content = generate_static_scope_section(company_name)
            gen_content = generate_section(full_prompt, token_limit, model=MODEL)
            # raw_content = static_content + "\n" + gen_content
            # content = normalize_newlines(raw_content)
            content = static_content + "\n" + gen_content
        else:
            section_prompt = prompt_values[i]
            full_prompt = build_prompt(global_prompt, minutes, section_prompt, token_limit)
            # raw_content = generate_section(full_prompt, token_limit, model=MODEL)
            # content = normalize_newlines(raw_content)

            content = generate_section(full_prompt, token_limit, model=MODEL)
            # Add in extra new line
            content = "\n" + content

            if heading == "Conclusion":
                content = content + "\n"

        # Add styled heading
        if heading in BM_SECTIONS:
            # Insert "Business Model" heading once
            if not inserted_bm_heading:
                bm_para = doc.add_paragraph("Business Model", style='Heading 1')
                bm_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                bm_run = bm_para.runs[0]
                bm_run.font.name = 'Calibri'
                bm_run.font.size = Pt(34)
                bm_run.font.color.rgb = RGBColor(255, 153, 0)
                bm_run.bold = True
                inserted_bm_heading = True

            # Create unstyled heading (NOT "Heading 1") for BM section
            heading_para = doc.add_paragraph()
            run = heading_para.add_run(heading)
            run.font.name = 'Calibri'
            run.font.size = Pt(34)
            run.font.color.rgb = RGBColor(255, 153, 0)
            run.bold = True

        else:
            # Styled heading that WILL appear in the table of contents
            heading_para = doc.add_paragraph(heading, style='Heading 1')
            run = heading_para.runs[0]
            run.font.name = 'Calibri'
            run.font.size = Pt(34)
            run.font.color.rgb = RGBColor(255, 153, 0)
            run.bold = True

        # Add normal body text
        # doc.add_paragraph(content)
        # Should do bolding AND bullet points
        # save_raw_text(heading, content)

        for line in content.split("\n"):
            stripped = line.strip()

            if not stripped:
                doc.add_paragraph()
                continue

            if is_bullet_point(stripped):
                # Strip hyphen/bullet prefix
                bullet_text = re.sub(r"^[-–—•●]\s+", "", stripped)
                # Handle markdown-style bold within the bullet
                add_markdown_bold_paragraph(doc, bullet_text, style="List Bullet")
            else:
                add_markdown_bold_paragraph(doc, stripped)

        if i != len(sections) - 1:
            doc.add_page_break()

    insert_logo(doc, "Logo3.png")
    # Add "Momentum Mind Lab Team" below the logo
    team_para = doc.add_paragraph()
    team_run = team_para.add_run("\nMomentum Mind Lab Team")
    team_run.font.name = 'Calibri'
    team_run.font.size = Pt(12)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)  # Move back to the beginning so Streamlit can read it
    return buffer

# Call OpenAI API to generate a section
def generate_section(full_prompt, token_limit, model=MODEL):
    response = openai.ChatCompletion.create(
        model=model,
        messages=[{"role": "user", "content": full_prompt}],
        max_tokens=int(token_limit * 1.3),  # 30% buffer
        temperature=0.7  # Slight randomness, can adjust
    )
    return response["choices"][0]["message"]["content"]

# Optional: Generate all sections in order (if needed later)
def generate_all_sections(global_prompt, minutes, prompt_library, sections, model=MODEL):
    results = []

    prompt_values = list(prompt_library.values())  # Rely on index order

    for i, (heading, token_limit) in enumerate(sections):
        section_prompt = prompt_values[i]
        full_prompt = build_prompt(global_prompt, minutes, section_prompt, token_limit)
        section_text = generate_section(full_prompt, token_limit, model=model)
        results.append((heading, section_text))

    return results  # List of (heading, generated_text)

def read_minutes(file_path):
    doc = Document(file_path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

# Shitty Wrapper Function (I <3 Overhead)
def generate_strategy_docx(minutes, file_path, company_name, status_area=None) -> BytesIO:
    with open("prompts.json", "r", encoding="utf-8") as f:
        prompts = json.load(f)

    GLOBAL_PROMPT = build_global(company_name)

    buffer = write_to_docx(
    file_path=file_path,
    global_prompt=GLOBAL_PROMPT,
    minutes=minutes,
    prompt_library=prompts,
    sections=SECTIONS,
    company_name=company_name,
    status_area=status_area)

    return buffer