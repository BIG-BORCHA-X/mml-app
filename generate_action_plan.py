import json
import re
from datetime import datetime, timedelta
from io import BytesIO

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement, ns

def add_markdown_bold_paragraph(doc, text, style="Normal"):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(0)

    # Indent bullets only
    if style == "List Bullet":
        paragraph.paragraph_format.left_indent = Inches(0.5)

    # Split into parts by bold markers (**...**)
    parts = re.split(r"(\*\*.*?\*\*)", text)

    for part in parts:
        run = paragraph.add_run()
        run.font.name = 'Calibri'
        run.font.size = Pt(12)

        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
        else:
            run.text = part

    return paragraph

def read_minutes(file_path):
    doc = Document(file_path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

def extract_json_from_response(content):
    match = re.search(r"\[\s*\{.*?\}\s*\]", content, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(0))
        except json.JSONDecodeError:
            print("❌ JSON format was still invalid.")
    print("❌ Could not extract valid JSON.")
    return []

def set_landscape_a4(doc):
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    margin = Inches(1)
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin

def set_column_width(cell, width_inches):
    cell.width = Inches(width_inches)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(ns.qn('w:w'), str(int(width_inches * 1440)))
    tcW.set(ns.qn('w:type'), 'dxa')
    tcPr.append(tcW)
    
def set_column_width2(table):
    widths = (Inches(1), Inches(2), Inches(1.5))
    widths = (Inches(0.68), Inches(1.15), Inches(1.7), Inches(3.1), Inches(1.17), Inches(1.77))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

def set_cell_margins(cell, top=102, start=102, bottom=102, end=102):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for name, value in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
        mar = OxmlElement(f'w:{name}')
        mar.set(ns.qn('w:w'), str(value))
        mar.set(ns.qn('w:type'), 'dxa')
        tcMar.append(mar)
    tcPr.append(tcMar)

def get_day_suffix(day):
    if 11 <= day <= 13:
        return "th"
    last_digit = day % 10
    return {1: "st", 2: "nd", 3: "rd"}.get(last_digit, "th")

def convert_when_to_date(_):
    startby = datetime.today() + timedelta(days=2)
    target = datetime.today() + timedelta(weeks=4)
    day1 = startby.day
    suffix1 = get_day_suffix(day1)
    day2 = target.day
    suffix2 = get_day_suffix(day2)
    formatted_date = f"Start {startby.strftime('%B')} {day1}{suffix1}, Complete by {target.strftime('%B')} {day2}{suffix2}"
    return formatted_date

def write_action_plan_docx(file_path, action_plan) -> BytesIO:
    doc = Document()
    set_landscape_a4(doc)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_para.add_run("Action Plan")
    run.font.name = 'Calibri'
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 153, 0)  # Orange

    # Table
    headers = ["Priority", "What", "Why", "How", "When", "Success Criteria"]
    raw_col_widths = [1.72, 3.62, 5.24, 6.27, 3.28, 4.37]
    col_widths = [x/2.54 for x in raw_col_widths]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        # set_column_width(cell, col_widths[i])
        # set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.bold = True

    # # Then set header row height
    # header_row = table.rows[0]._tr
    # trPr = header_row.get_or_add_trPr()
    # trHeight = OxmlElement('w:trHeight')
    # trHeight.set(ns.qn('w:val'), '518')
    # trHeight.set(ns.qn('w:hRule'), 'exact')
    # trPr.append(trHeight)

    # Priority Emojies
    priority_map = {
        "Red": "🔴",
        "Yellow": "🟡",
        "Green": "🟢"
    }

    # Action Plan rows
    for idx, row in enumerate(action_plan, start=1):
        row["When"] = convert_when_to_date(row["When"])
        row_cells = table.add_row().cells
        for i, key in enumerate(headers):
            cell = row_cells[i]
            set_cell_margins(cell)
            
            # Handle bullet points in HOW
            if key == "How" and isinstance(row[key], list):
                cell._element.clear_content()
                for bullet in row[key]:
                    bullet_para = cell.add_paragraph()
                    # bullet_para.paragraph_format.left_indent = Inches(0.2)  # Optional indent
                    # bullet_run = bullet_para.add_run(f"• {bullet}")
                    bullet_run = bullet_para.add_run(f"- {bullet}")
                    bullet_run.font.name = 'Calibri'
                    bullet_run.font.size = Pt(12)
            else:
                para = cell.paragraphs[0]
                value = str(row[key])

                # Add numbering to "What", add emojis to "Priority"
                # if key == "What":
                #     value = f"{idx}. {value}"
                # elif key == "Priority":
                #     value = priority_map[value]
                
                # Priority plus Numbering
                if key == "Priority":
                    value = f"{idx}. {priority_map[value]}"

                run = para.add_run(value)
                run.font.name = 'Calibri'
                run.font.size = Pt(12)
                if key == "What" or key == "Priority":
                    run.bold = True

    # Additional Notes
    # Notes about priority
    notes = [
        "**Key**:",
        "🔴 **High priority**: Critical for launch, client delivery, or business continuity",
        "🟡 **Medium priority**: Important but not immediately time-sensitive",
        "🟢 **Low priority**: Valuable for long-term improvements or future planning"
    ]

    for line in notes:
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph()  # Maintain spacing
            continue

        add_markdown_bold_paragraph(doc, stripped)

    # Set Column Width, needs to be performed on ALL cells in grid.
    for column in range(len(col_widths)):
        for row in table.columns[column].cells:
            row.width = Inches(col_widths[column])
            set_cell_margins(row)

    # Save file
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)  # Move back to the beginning so Streamlit can read it
    return buffer